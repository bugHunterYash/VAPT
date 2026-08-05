import docx
import os

def clean_template():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(base_dir, "master_template.docx")
    output_path = os.path.join(base_dir, "clean_template.docx")
    
    doc = docx.Document(input_path)
    
    chart_intro_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "The graph below shows a summary of the total number of vulnerabilities" in p.text:
            chart_intro_idx = i
            break
            
    if chart_intro_idx != -1:
        # 1. First, make sure {{ chart_image }} is there (if not already added by previous run)
        p_intro = doc.paragraphs[chart_intro_idx]
        if "{{ chart_image }}" not in p_intro.text and "{{ chart_image }}" not in doc.paragraphs[chart_intro_idx+1].text:
            new_p = p_intro.insert_paragraph_before("{{ chart_image }}")
            p_intro._p.addnext(new_p._p)
            
        paragraphs_to_delete = []
        for j in range(chart_intro_idx, len(doc.paragraphs)):
            p_curr = doc.paragraphs[j]
            if "3.2" in p_curr.text and "Scope" in p_curr.text:
                break
                
            # If the paragraph is empty now and not the intro or the chart_image placeholder, remove it
            if j > chart_intro_idx and not p_curr.text.strip() and "{{ chart_image }}" not in p_curr.text:
                paragraphs_to_delete.append(p_curr)
                
        for p in paragraphs_to_delete:
            p._element.getparent().remove(p._element)
                
    # 2. Find the POC endfor and add finding endfor (handled previously, but let's redo cleanly from master)
    # Actually, we should always operate on master_template.docx
    inner_endfor_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "{%p endfor %}" in p.text:
            inner_endfor_idx = i
            break
            
    if inner_endfor_idx != -1:
        # Check if we already have two consecutive endfors
        if "{%p endfor %}" not in doc.paragraphs[inner_endfor_idx+1].text:
            p_inner_end = doc.paragraphs[inner_endfor_idx]
            new_end_p = p_inner_end.insert_paragraph_before("{%p endfor %}")
            p_inner_end._p.addnext(new_end_p._p)
        
    doc.save(output_path)
    
    doc = docx.Document(output_path)
    
    finding_endfor_idx = -1
    conclusion_idx = -1
    
    found_first_endfor = False
    for i, p in enumerate(doc.paragraphs):
        if "{%p endfor %}" in p.text:
            if not found_first_endfor:
                found_first_endfor = True
            else:
                finding_endfor_idx = i
                break
                
    for i in range(finding_endfor_idx + 1, len(doc.paragraphs)):
        if "5. Conclusion" in doc.paragraphs[i].text:
            conclusion_idx = i
            break
            
    if finding_endfor_idx != -1 and conclusion_idx != -1:
        body_elements = list(doc.part.element.body)
        finding_end_elem = doc.paragraphs[finding_endfor_idx]._p
        conclusion_elem = doc.paragraphs[conclusion_idx]._p
        
        in_delete_zone = False
        elements_to_delete = []
        
        for elem in body_elements:
            if elem == finding_end_elem:
                in_delete_zone = True
                continue
            if elem == conclusion_elem:
                in_delete_zone = False
                break
            if in_delete_zone:
                elements_to_delete.append(elem)
                
        for elem in elements_to_delete:
            elem.getparent().remove(elem)
            
    # 5. Delete trailing {%p endfor %}
    for p in reversed(doc.paragraphs):
        if "{%p endfor %}" in p.text:
            p._element.getparent().remove(p._element)
            break
            
    # 6. Clean Table Data Rows
    def clear_table_rows(table):
        while len(table.rows) > 1:
            tr = table.rows[-1]._tr
            tr.getparent().remove(tr)

    for t in doc.tables:
        if len(t.rows) > 0 and len(t.columns) == 4 and "S. No." in t.cell(0,0).text and "Date of Revision" in t.cell(0,1).text:
            clear_table_rows(t)
        elif len(t.rows) > 0 and len(t.columns) == 5 and "S. No" in t.cell(0,0).text and "Designation" in t.cell(0,2).text:
            clear_table_rows(t)
        elif len(t.rows) > 0 and len(t.columns) == 3 and "Application Name" in t.cell(0,1).text:
            clear_table_rows(t)
        elif len(t.rows) > 0 and len(t.columns) == 3 and "Application Username" in t.cell(0,1).text:
            clear_table_rows(t)
        elif len(t.rows) > 0 and len(t.columns) == 4 and "S. No." in t.cell(0,0).text and "Count of Vulnerabilities" in t.cell(0,3).text:
            clear_table_rows(t)

    doc.save(output_path)
    print("Template cleaned successfully: clean_template.docx")

if __name__ == "__main__":
    clean_template()
