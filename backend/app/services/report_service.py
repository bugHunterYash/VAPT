import io
import json
import datetime
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from app.services.storage_service import get_evidence_bytes

def add_evidence_image(doc, file_path, caption, max_width_inches=6.0):
    try:
        image_bytes = get_evidence_bytes(file_path)
        img_stream = io.BytesIO(image_bytes)
        doc.add_picture(img_stream, width=Inches(max_width_inches))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].font.italic = True
    except Exception as e:
        doc.add_paragraph(f"[Missing Image: {file_path}]")
        print(f"Error adding image {file_path}: {e}")

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

def generate_methodology_graphic() -> io.BytesIO:
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    
    fig, ax = plt.subplots(figsize=(10, 2.5))
    ax.axis('off')
    
    steps = ["Planning", "Discovery", "Attack", "Review &\nAnalysis", "Reporting"]
    colors = ['#5B9BD5', '#5B9BD5', '#5B9BD5', '#5B9BD5', '#5B9BD5']
    
    start_x = 0
    y = 0.5
    w = 1.6
    h = 0.8
    arrow_w = 0.4
    gap = 0.2
    
    for i, step in enumerate(steps):
        # Draw chevron
        if i == 0:
            # First one is a rectangle with arrow head
            poly = patches.Polygon([[start_x, y-h/2], [start_x+w, y-h/2], [start_x+w+arrow_w, y], [start_x+w, y+h/2], [start_x, y+h/2]], 
                                   closed=True, facecolor=colors[i], edgecolor='white', lw=2)
        else:
            # Others have arrow tail indent
            poly = patches.Polygon([[start_x, y-h/2], [start_x+w, y-h/2], [start_x+w+arrow_w, y], [start_x+w, y+h/2], [start_x, y+h/2], [start_x+arrow_w, y]], 
                                   closed=True, facecolor=colors[i], edgecolor='white', lw=2)
        
        # Add shadow
        import copy
        shadow = copy.copy(poly)
        shadow.set_facecolor('gray')
        shadow.set_alpha(0.3)
        shadow.set_xy(shadow.get_xy() + [0.03, -0.05])
        ax.add_patch(shadow)
        ax.add_patch(poly)
        
        # Text
        text_x = start_x + w/2 + (arrow_w/2 if i>0 else arrow_w/4)
        ax.text(text_x, y, step, ha='center', va='center', fontweight='bold', color='black', fontsize=11)
        
        start_x += w + gap
        
    # Draw curved arrow back
    style = "Simple, tail_width=3, head_width=10, head_length=12"
    kw = dict(arrowstyle=style, color="#5B9BD5")
    # From Review (index 3) to Discovery (index 1)
    x_start = (w + gap) * 3 + w/2
    x_end = (w + gap) * 1 + w/2
    a = patches.FancyArrowPatch((x_start, y - h/2 - 0.1), (x_end, y - h/2 - 0.1),
                                connectionstyle="arc3,rad=0.4", **kw)
    ax.add_patch(a)
    
    ax.set_xlim(-0.5, start_x + 0.5)
    ax.set_ylim(-0.5, 1.5)
    
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    img_stream.seek(0)
    return img_stream

def generate_severity_chart(counts: dict) -> io.BytesIO:
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    from matplotlib.patches import Patch
    import matplotlib.colors as mcolors
    import numpy as np
    
    labels = ['Critical', 'High', 'Medium', 'Low', 'Informative']
    values = [counts.get('Critical', 0), counts.get('High', 0), counts.get('Medium', 0), counts.get('Low', 0), counts.get('Info', counts.get('Informative', 0))]
    base_colors = ['#FF0000', '#FF9900', '#FFFF00', '#0070C0', '#BFBFBF']
    
    fig, ax = plt.subplots(figsize=(7, 4), facecolor='#E6F0FA')
    ax.set_facecolor('#E6F0FA')
    
    # Hide standard borders and ticks
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.set_xticks([])
    ax.set_yticks([])
    
    max_val = max(values) if values else 0
    y_max = max(max_val + 1, 2)
    
    # Grid lines (simulate back wall)
    for i in range(y_max + 1):
        ax.plot([-0.5, len(labels) - 0.5], [i, i], color='gray', linestyle='-', linewidth=0.5, alpha=0.3)
        ax.text(-0.6, i, str(i), va='center', ha='right', color='black', fontsize=10)
        
    width = 0.4
    dx = 0.12 # depth X
    dy = 0.12 * y_max / len(labels) * 1.5 # depth Y, scaled by aspect ratio roughly
    
    for i, (label, val, color) in enumerate(zip(labels, values, base_colors)):
        x_center = i
        x_left = x_center - width/2
        x_right = x_center + width/2
        
        # Determine shade colors
        c_rgb = mcolors.hex2color(color)
        c_dark = tuple(max(0, c * 0.7) for c in c_rgb)
        c_light = tuple(min(1, c * 1.2) for c in c_rgb)
        
        # 1. Front face
        front = patches.Rectangle((x_left, 0), width, val, facecolor=color, edgecolor='black', linewidth=0.5)
        ax.add_patch(front)
        
        # 2. Right face
        right_poly = patches.Polygon([
            [x_right, 0],
            [x_right + dx, dy],
            [x_right + dx, val + dy],
            [x_right, val]
        ], closed=True, facecolor=c_dark, edgecolor='black', linewidth=0.5)
        ax.add_patch(right_poly)
        
        # 3. Top face
        top_poly = patches.Polygon([
            [x_left, val],
            [x_right, val],
            [x_right + dx, val + dy],
            [x_left + dx, val + dy]
        ], closed=True, facecolor=c_light, edgecolor='black', linewidth=0.5)
        ax.add_patch(top_poly)
        
        # Label below bar
        ax.text(x_center, -0.3, label, ha='center', va='top', fontsize=10, color='black')
        
        # Data label above bar
        label_text = f"{label}, {val}"
        ax.text(x_center + dx/2, val + dy + (y_max * 0.05), label_text, 
                ha='center', va='bottom', fontsize=9, color='black',
                bbox=dict(facecolor='white', alpha=0.8, edgecolor='none', pad=2))

    ax.set_title('Vulnerabilities', fontweight='bold', color='black', size=16, pad=20)
    
    # Custom Legend
    legend_elements = [Patch(facecolor=base_colors[i], edgecolor='black', label=labels[i]) for i in range(len(labels))]
    ax.legend(handles=legend_elements, loc='lower center', bbox_to_anchor=(0.5, -0.2), ncol=5, frameon=False, handletextpad=0.2, columnspacing=1)
    
    ax.set_xlim(-1, len(labels))
    ax.set_ylim(-0.5, y_max + (y_max * 0.15))
    
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=300, facecolor=fig.get_facecolor(), edgecolor='none', bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf

def apply_heading_style(doc, level, text):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 176, 240) # Professional Blue
    return h

def generate_docx_report(project_data: dict) -> io.BytesIO:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Inches
    import docx
    import os
    import datetime
    import io
    import json
    
    template_path = os.path.join(os.path.dirname(__file__), "clean_template.docx")
    
    meta = project_data.get('reportMeta', {})
    if not meta: meta = {}
    
    # Extract Assessment Parameters
    assessments = project_data.get('assessmentParameters', [])
    if not assessments:
        # Fallback to single app if missing
        assessments = [{
            'applicationName': project_data.get('applicationName', project_data.get('name', 'N/A')),
            'applicationUrl': project_data.get('applicationUrl', 'N/A'),
            'username': project_data.get('username', 'N/A'),
            'password': project_data.get('password', 'N/A')
        }]
    
    findings_raw = project_data.get('findings', [])
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Info": 0}
    
    # 1. Pre-process the docx tables manually before docxtpl rendering
    doc = docx.Document(template_path)
    
    for t in doc.tables:
        if len(t.rows) > 0 and len(t.columns) == 4 and "S. No." in t.cell(0,0).text and "Date of Revision" in t.cell(0,1).text:
            revs = meta.get('revisionHistory', [])
            if not revs:
                revs = [{'date': datetime.datetime.now().strftime("%d/%m/%Y"), 'version': '1.0', 'description': 'Initial Release'}]
            for i, rev in enumerate(revs, 1):
                new_row = t.add_row()
                new_row.cells[0].text = str(i)
                new_row.cells[1].text = str(rev.get('date', ''))
                new_row.cells[2].text = str(rev.get('version', ''))
                new_row.cells[3].text = str(rev.get('description', ''))
                
        elif len(t.rows) > 0 and len(t.columns) == 5 and "S. No" in t.cell(0,0).text and "Designation" in t.cell(0,2).text:
            team = project_data.get('teamMembers', [])
            for i, member in enumerate(team, 1):
                new_row = t.add_row()
                new_row.cells[0].text = str(i)
                new_row.cells[1].text = str(member.get('name', ''))
                new_row.cells[2].text = str(member.get('designation', ''))
                new_row.cells[3].text = str(member.get('email', ''))
                new_row.cells[4].text = str(member.get('qualifications', ''))
                
        elif len(t.rows) > 0 and len(t.columns) == 3 and "Application Name" in t.cell(0,1).text:
            for i, asm in enumerate(assessments, 1):
                new_row = t.add_row()
                new_row.cells[0].text = str(i)
                new_row.cells[1].text = str(asm.get('applicationName', ''))
                new_row.cells[2].text = str(asm.get('applicationUrl', ''))
                
        elif len(t.rows) > 0 and len(t.columns) == 3 and "Application Username" in t.cell(0,1).text:
            if meta.get('includeCredentials', False):
                for i, asm in enumerate(assessments, 1):
                    new_row = t.add_row()
                    new_row.cells[0].text = str(i)
                    new_row.cells[1].text = str(asm.get('username', ''))
                    new_row.cells[2].text = str(asm.get('password', ''))
            else:
                # Clear entire table if credentials not included
                tbl = t._element
                tbl.getparent().remove(tbl)

        elif len(t.rows) > 0 and len(t.columns) == 4 and "S. No." in t.cell(0,0).text and "Count of Vulnerabilities" in t.cell(0,3).text:
            for idx, f in enumerate(findings_raw, 1):
                new_row = t.add_row()
                new_row.cells[0].text = str(idx)
                new_row.cells[1].text = f.get('title', 'Unknown Vulnerability')
                sev = f.get('severity', 'Info')
                if sev == 'Informative': sev = 'Info'
                new_row.cells[2].text = sev
                new_row.cells[3].text = "1"
    
    # Fix split Jinja tags across multiple runs
    def consolidate_tags(paragraphs):
        for p in paragraphs:
            if "{{" in p.text and "}}" in p.text:
                full_text = p.text
                if len(p.runs) > 1:
                    for i, r in enumerate(p.runs):
                        if i == 0:
                            r.text = full_text
                        else:
                            r.text = ""
                            
    consolidate_tags(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                consolidate_tags(cell.paragraphs)
                
    temp_buffer = io.BytesIO()
    doc.save(temp_buffer)
    temp_buffer.seek(0)
    
    tpl = DocxTemplate(temp_buffer)
    
    findings = []
    global_ev_counter = 1
    
    for idx, f in enumerate(findings_raw, 1):
        sev = f.get('severity', 'Info')
        if sev == 'Informative': sev = 'Info'
        counts[sev] = counts.get(sev, 0) + 1
        
        param_val = f.get('parameter', '')
        parsed_params = "N/A"
        try:
            pl = json.loads(param_val)
            if isinstance(pl, list) and len(pl) > 0:
                parsed_params = "\n".join([f"{p.get('name')} = {p.get('value')}" if p.get('value') else p.get('name') for p in pl])
            else:
                parsed_params = param_val or "N/A"
        except:
            parsed_params = param_val or "N/A"
            
        # Build Subdoc for bullets
        impact_sub = tpl.new_subdoc()
        impact_raw = f.get('impact', '')
        if isinstance(impact_raw, list):
            for i, imp in enumerate(impact_raw):
                impact_sub.add_paragraph(imp, style='List Paragraph' if i > 0 else None)
        else:
            impact_sub.add_paragraph(impact_raw)
            
        mit_sub = tpl.new_subdoc()
        mit_raw = f.get('mitigation', '')
        if isinstance(mit_raw, list):
            for i, mit in enumerate(mit_raw):
                mit_sub.add_paragraph(mit, style='List Paragraph' if i > 0 else None)
        else:
            mit_sub.add_paragraph(mit_raw)
            
        url_val = f.get('affectedUrls', '')
        if url_val: url_val = "\n".join([u.strip() for u in url_val.split('\n') if u.strip()])
        
        evidences = []
        for ev in f.get('evidences', []):
            if ev.get('filePath'):
                try:
                    img_bytes = get_evidence_bytes(ev.get('filePath'))
                    img_stream = io.BytesIO(img_bytes)
                    img = InlineImage(tpl, img_stream, width=Inches(6.0))
                    evidences.append({
                        'image': img, 
                        'caption': ev.get('caption', ''),
                        'global_index': global_ev_counter
                    })
                    global_ev_counter += 1
                except Exception as e:
                    print(f"Error loading image: {e}")
        
        finding_data = {
            'index': idx,
            'title': f.get('title', 'Unknown Vulnerability'),
            'owasp': f.get('owasp', 'N/A'),
            'cwe': f.get('cwe', 'N/A'),
            'severity': f.get('severity', 'N/A'),
            'description': f.get('description', 'N/A'),
            'url': url_val or "N/A",
            'parameter': parsed_params,
            'impact': impact_sub,
            'mitigation': mit_sub,
            'evidences': evidences
        }
        findings.append(finding_data)
        
    chart_buf = generate_severity_chart(counts)
    chart_image = InlineImage(tpl, chart_buf, width=Inches(6.5))
    
    # Add date formatting helper
    def format_date(d_str):
        if not d_str: return datetime.datetime.now().strftime("%d/%m/%Y")
        try:
            if 'T' in d_str:
                dt = datetime.datetime.fromisoformat(d_str.replace('Z', '+00:00'))
                return dt.strftime("%d/%m/%Y")
        except Exception:
            pass
        return d_str

    context = {
        'reportTitle': meta.get('documentName', 'Web Application Detailed Vulnerabilities Report'),
        'documentVersion': meta.get('documentVersion', '1.0'),
        'preparedBy': meta.get('preparedBy', 'VAPT Team'),
        'approvedBy': meta.get('approvedBy', ''),
        'reviewedBy': meta.get('reviewedBy', ''),
        'releasedBy': meta.get('releasedBy', ''),
        'certInEmpanelment': meta.get('certInEmpanelment', ''),
        'clientName': meta.get('clientName', 'Organization'),
        'documentClassification': meta.get('documentClassification', 'CONFIDENTIAL'),
        'chart_image': chart_image,
        
        'applicationName': assessments[0]['applicationName'] if assessments else 'N/A',
        'organization': meta.get('organization', 'Organization'),
        'documentDate': format_date(meta.get('documentDate')),
        'startDate': format_date(meta.get('startDate')),
        'totalVulns': sum(counts.values()),
        'critVulns': counts['Critical'],
        'highVulns': counts['High'],
        'medVulns': counts['Medium'],
        'lowVulns': counts['Low'],
        'infoVulns': counts['Info'],
        'findings': findings
    }
    
    # Generation-time Integrity Checks
    expected_findings_count = len(findings_raw)
    calculated_findings_count = sum(counts.values())
    if expected_findings_count != calculated_findings_count:
        raise ValueError(f"Integrity Check Failed: Expected {expected_findings_count} total severity count, but calculated {calculated_findings_count}.")
        
    if len(findings) != expected_findings_count:
        raise ValueError(f"Integrity Check Failed: Expected {expected_findings_count} generated findings, but prepared {len(findings)}.")
        
    tpl.render(context)
    
    # Post-render Image Blob Replacement
    cover_image_b64 = project_data.get('coverImage')
    thank_you_image_b64 = project_data.get('thankYouImage')
    
    if cover_image_b64 or thank_you_image_b64:
        from PIL import Image
        import base64
        
        def replace_blob_with_crop(part, b64_str):
            if ',' in b64_str:
                b64_str = b64_str.split(',')[1]
            try:
                new_bytes = base64.b64decode(b64_str)
                orig_img = Image.open(io.BytesIO(part.blob))
                orig_w, orig_h = orig_img.size
                orig_aspect = orig_w / orig_h
                
                new_img = Image.open(io.BytesIO(new_bytes))
                new_w, new_h = new_img.size
                new_aspect = new_w / new_h
                
                if abs(orig_aspect - new_aspect) > 0.05:
                    if new_aspect > orig_aspect:
                        # Wider: crop sides
                        target_w = int(new_h * orig_aspect)
                        left = (new_w - target_w) // 2
                        new_img = new_img.crop((left, 0, left + target_w, new_h))
                    else:
                        # Taller: crop top/bottom
                        target_h = int(new_w / orig_aspect)
                        top = (new_h - target_h) // 2
                        new_img = new_img.crop((0, top, new_w, top + target_h))
                
                out_buf = io.BytesIO()
                new_img.save(out_buf, format=orig_img.format or 'PNG')
                part._blob = out_buf.getvalue()
            except Exception as e:
                print(f"Error replacing image: {e}")

        # Find cover image (first image in document)
        if cover_image_b64:
            for p in tpl.docx.paragraphs:
                blips = p._p.xpath('.//a:blip')
                if blips:
                    rId = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    part = tpl.docx.part.rels[rId].target_part
                    replace_blob_with_crop(part, cover_image_b64)
                    break
                    
        # Find thank you image (last image in document)
        if thank_you_image_b64:
            for p in reversed(tpl.docx.paragraphs):
                blips = p._p.xpath('.//a:blip')
                if blips:
                    rId = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    part = tpl.docx.part.rels[rId].target_part
                    replace_blob_with_crop(part, thank_you_image_b64)
                    break

    buffer = io.BytesIO()
    tpl.save(buffer)
    buffer.seek(0)
    
    # Final Validation: Scan for unresolved tags
    final_doc = docx.Document(buffer)
    for p in final_doc.paragraphs:
        if "{{" in p.text or "}}" in p.text:
            raise ValueError(f"Integrity Check Failed: Unresolved template syntax found in paragraph: '{p.text}'")
    for table in final_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{" in p.text or "}}" in p.text:
                        raise ValueError(f"Integrity Check Failed: Unresolved template syntax found in table cell: '{p.text}'")
                        
    buffer.seek(0)
    return buffer


import shutil

def find_libreoffice():
    # 0. Check Environment Variable
    env_path = os.environ.get("LIBREOFFICE_PATH")
    if env_path and os.path.exists(env_path): return env_path
    
    # 1. Check shutil.which("soffice")
    path = shutil.which("soffice")
    if path: return path
    # 2. Check shutil.which("libreoffice")
    path = shutil.which("libreoffice")
    if path: return path
    # 3. Explicit check: C:\Program Files\LibreOffice\program\soffice.exe
    p1 = r"C:\Program Files\LibreOffice\program\soffice.exe"
    if os.path.exists(p1): return p1
    # 4. Explicit check: C:\Program Files (x86)\LibreOffice\program\soffice.exe
    p2 = r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
    if os.path.exists(p2): return p2
    return None

def generate_pdf_report(project_data: dict) -> io.BytesIO:
    print("[REPORT] Requested format: PDF")
    docx_buffer = generate_docx_report(project_data)
    
    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    with os.fdopen(fd, 'wb') as f:
        f.write(docx_buffer.getvalue())
        
    outdir = os.path.dirname(docx_path)
    pdf_path = docx_path.replace(".docx", ".pdf")
    
    lo_path = find_libreoffice()
    if not lo_path:
        os.remove(docx_path)
        raise Exception("LibreOffice executable not found. Cannot convert DOCX to PDF.")
        
    print(f"[REPORT] DOCX generated: {docx_path}")
    print(f"[PDF] LibreOffice detected: {lo_path}")
    print(f"[PDF] Output directory: {outdir}")
    
    cmd = [
        lo_path,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        outdir,
        docx_path
    ]
    
    try:
        print(f"[PDF] Executing command: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True)
        print(f"[PDF] Return code: {result.returncode}")
        print(f"[PDF] stdout: {result.stdout}")
        print(f"[PDF] stderr: {result.stderr}")
        
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed with return code {result.returncode}.\nStdout: {result.stdout}\nStderr: {result.stderr}")
            
    except Exception as e:
        if os.path.exists(docx_path): os.remove(docx_path)
        raise Exception(f"PDF conversion process failed: {str(e)}")
        
    if not os.path.exists(pdf_path):
        os.remove(docx_path)
        raise Exception(f"PDF conversion failed: Expected PDF file does not exist at {pdf_path}. LibreOffice may have failed silently. Check stderr.")
        
    if os.path.getsize(pdf_path) == 0:
        os.remove(docx_path)
        os.remove(pdf_path)
        raise Exception(f"PDF conversion failed: Output PDF file is empty.")
        
    with open(pdf_path, 'rb') as f:
        pdf_bytes = f.read()
        
    if os.path.exists(docx_path): os.remove(docx_path)
    if os.path.exists(pdf_path): os.remove(pdf_path)
    
    return io.BytesIO(pdf_bytes)

def generate_excel_report(project_data: dict) -> io.BytesIO:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Vulnerability Tracker"
    
    headers = [
        "S. No", "Vulnerability Name", "Owasp Mapping", "Severity", 
        "Vulnerable URL", "Affected Parameters", "Solution"
    ]
    ws.append(headers)
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'), 
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align
        cell.border = border
        
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 45
    ws.column_dimensions['F'].width = 23
    ws.column_dimensions['G'].width = 120
        
    findings = project_data.get("findings", [])
    
    for idx, finding in enumerate(findings, 2):
        ws.row_dimensions[idx].height = 80
        
        c_sno = ws.cell(row=idx, column=1, value=idx-1)
        c_sno.border = border
        c_sno.alignment = center_align
        
        c_title = ws.cell(row=idx, column=2, value=finding.get('title', ''))
        c_title.border = border
        c_title.alignment = center_align
        
        c_owasp = ws.cell(row=idx, column=3, value=finding.get('owasp', ''))
        c_owasp.border = border
        c_owasp.alignment = center_align
        
        sev = finding.get('severity', '')
        c_sev = ws.cell(row=idx, column=4, value=sev)
        c_sev.border = border
        c_sev.alignment = center_align
        
        if sev == "Critical":
            c_sev.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
            c_sev.font = Font(color="FFFFFF")
        elif sev == "High":
            c_sev.fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
        elif sev == "Medium":
            c_sev.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        elif sev == "Low":
            c_sev.fill = PatternFill(start_color="0070C0", end_color="0070C0", fill_type="solid")
            c_sev.font = Font(color="FFFFFF")
        elif sev == "Info" or sev == "Informative":
            c_sev.fill = PatternFill(start_color="BFBFBF", end_color="BFBFBF", fill_type="solid")
            
        urls_raw = finding.get('affectedUrls', '')
        url_list = [u.strip() for u in urls_raw.split('\n') if u.strip()] if urls_raw else []
        c_url = ws.cell(row=idx, column=5, value="\n".join(url_list) if url_list else "N/A")
        c_url.border = border
        c_url.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        if len(url_list) == 1 and url_list[0].startswith("http"):
            c_url.hyperlink = url_list[0]
            c_url.font = Font(color="0563C1", underline="single")
            
        param_val = finding.get('parameter', '')
        parsed_params = "N/A"
        if param_val:
            try:
                param_list = json.loads(param_val)
                if isinstance(param_list, list) and len(param_list) > 0:
                    lines = []
                    for p in param_list:
                        name = p.get('name', '')
                        value = p.get('value', '')
                        if value: lines.append(f"{name} = {value}")
                        else: lines.append(name)
                    parsed_params = "\n".join(lines)
                else:
                    parsed_params = str(param_val)
            except Exception:
                parsed_params = str(param_val)
                
        c_param = ws.cell(row=idx, column=6, value=parsed_params)
        c_param.border = border
        c_param.alignment = center_align
        
        mitigation_data = finding.get('mitigation', '')
        if isinstance(mitigation_data, list):
            mitigation_val = "\n\n".join(f"• {m}" for m in mitigation_data if m)
        else:
            mitigation_val = str(mitigation_data)
            
        c_sol = ws.cell(row=idx, column=7, value=mitigation_val)
        c_sol.border = border
        c_sol.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    
    ws.print_options.horizontalCentered = True
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.print_title_rows = '1:1'
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
